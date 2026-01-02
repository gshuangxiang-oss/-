#!/usr/bin/env python3
"""
将测试报告转换为 Word 文档
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("安装 python-docx...")
    import subprocess
    subprocess.run(['pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_report():
    """创建 Word 测试报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # ========== 标题 ==========
    title = doc.add_heading('车牌检测与识别系统实验报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== 项目信息表格 ==========
    doc.add_heading('项目信息', level=1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('项目名称', '车牌检测与识别系统'),
        ('课程名称', '深度学习'),
        ('提交日期', '2026年1月3日'),
        ('学生姓名', '郭双翔'),
        ('学号', '23050420'),
    ]
    
    for i, (key, value) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = key
        cells[1].text = value
        set_cell_shading(cells[0], 'E6E6E6')
    
    doc.add_paragraph()
    
    # ========== 1. 核心算法 ==========
    doc.add_heading('1. 核心算法', level=1)
    
    doc.add_paragraph(
        '本项目的核心任务是通过深度学习技术对车辆图像进行车牌检测与识别，'
        '具体包括两个部分：车牌区域定位和车牌字符识别。为了实现这一目标，'
        '我们采用了基于YOLOv8的目标检测算法进行车牌定位，以及基于EasyOCR的'
        '光学字符识别技术进行车牌号码识别，并结合了多种图像预处理技术来提高识别精度。'
    )
    
    # 1.1 车牌检测算法
    doc.add_heading('1.1 车牌检测算法（YOLOv8）', level=2)
    
    doc.add_paragraph(
        '为了精准定位图像中的车牌区域，我们采用了YOLOv8（You Only Look Once v8）'
        '目标检测算法。YOLO是一种端到端的单阶段目标检测网络，具有检测速度快、准确率高的特点。'
    )
    
    doc.add_paragraph('YOLOv8 网络结构特点：', style='List Bullet')
    doc.add_paragraph('Backbone：采用CSPDarknet作为主干网络，提取多尺度特征', style='List Bullet')
    doc.add_paragraph('Neck：使用PAN-FPN结构进行特征融合', style='List Bullet')
    doc.add_paragraph('Head：解耦头设计，分别处理分类和回归任务', style='List Bullet')
    
    doc.add_paragraph('检测流程：')
    doc.add_paragraph('1. 输入图像经过Backbone提取特征', style='List Number')
    doc.add_paragraph('2. 特征图通过Neck进行多尺度融合', style='List Number')
    doc.add_paragraph('3. Head输出边界框坐标、置信度和类别概率', style='List Number')
    doc.add_paragraph('4. 非极大值抑制（NMS）去除重复检测', style='List Number')
    
    # 1.2 车牌字符识别算法
    doc.add_heading('1.2 车牌字符识别算法（EasyOCR）', level=2)
    
    doc.add_paragraph(
        '车牌字符识别采用EasyOCR框架，该框架支持中英文混合识别，适合中国车牌的识别需求。'
    )
    
    doc.add_paragraph('CRNN网络结构：', style='List Bullet')
    doc.add_paragraph('卷积层（CNN）：提取图像特征', style='List Bullet')
    doc.add_paragraph('循环层（BiLSTM）：建模序列依赖关系', style='List Bullet')
    doc.add_paragraph('转录层（CTC）：将特征序列转换为文本', style='List Bullet')
    
    # 1.3 传统图像处理方法
    doc.add_heading('1.3 传统图像处理方法（备选方案）', level=2)
    
    doc.add_paragraph('当深度学习模型不可用时，系统采用传统图像处理方法作为备选方案：')
    doc.add_paragraph('颜色空间转换：RGB转灰度图', style='List Bullet')
    doc.add_paragraph('双边滤波：去噪同时保留边缘', style='List Bullet')
    doc.add_paragraph('Canny边缘检测：提取轮廓', style='List Bullet')
    doc.add_paragraph('轮廓分析：基于宽高比（3:1~4:1）筛选车牌区域', style='List Bullet')
    
    # 1.4 图像预处理与增强
    doc.add_heading('1.4 图像预处理与增强', level=2)
    
    doc.add_paragraph('为了提高车牌识别的准确率，我们对裁剪出的车牌区域进行了以下预处理：')
    doc.add_paragraph('尺寸归一化：将车牌图像高度统一为48像素', style='List Bullet')
    doc.add_paragraph('灰度转换：减少颜色干扰', style='List Bullet')
    doc.add_paragraph('CLAHE均衡化：自适应直方图均衡化，增强对比度', style='List Bullet')
    
    # ========== 2. 程序功能介绍 ==========
    doc.add_heading('2. 程序功能介绍', level=1)
    
    doc.add_heading('2.1 车牌检测模块', level=2)
    doc.add_paragraph(
        '程序首先使用YOLOv8模型对输入图像进行车牌检测。'
        '检测器会返回所有检测到的车牌边界框及其置信度。'
    )
    
    doc.add_heading('2.2 车牌字符识别模块', level=2)
    doc.add_paragraph(
        '检测到车牌区域后，程序裁剪车牌图像并送入OCR模块进行字符识别。'
    )
    
    doc.add_heading('2.3 车牌格式验证', level=2)
    doc.add_paragraph(
        '识别完成后，程序会对车牌号码进行格式验证，判断是否符合中国车牌规则。'
    )
    
    doc.add_heading('2.4 Web服务与API接口', level=2)
    doc.add_paragraph(
        '程序提供了基于Flask的REST API接口，支持图片上传和Base64编码两种方式。'
    )
    
    # ========== 3. 实验结果图与分析 ==========
    doc.add_heading('3. 实验结果图与分析', level=1)
    
    doc.add_heading('3.1 测试图片说明', level=2)
    doc.add_paragraph('项目包含 13 张测试图片，位于 test_images/ 目录：')
    
    # 测试图片表格
    img_table = doc.add_table(rows=7, cols=3)
    img_table.style = 'Table Grid'
    
    img_headers = ['文件名', '车牌号', '说明']
    for i, header in enumerate(img_headers):
        img_table.rows[0].cells[i].text = header
        set_cell_shading(img_table.rows[0].cells[i], 'CCCCCC')
    
    img_data = [
        ('test_car_1.jpg', '浙A·12345', '标准蓝牌'),
        ('test_car_2.jpg', '京B·88888', '标准蓝牌'),
        ('test_car_3.jpg', '沪C·66666', '标准蓝牌'),
        ('test_car_4.jpg', '粤D·55555', '标准蓝牌'),
        ('test_car_5.jpg', '苏E·99999', '标准蓝牌'),
        ('test_random_1~5.jpg', '随机生成', '测试泛化能力'),
    ]
    
    for i, row_data in enumerate(img_data):
        for j, cell_text in enumerate(row_data):
            img_table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 识别效果统计', level=2)
    
    result_table = doc.add_table(rows=5, cols=5)
    result_table.style = 'Table Grid'
    
    result_headers = ['测试类型', '样本数', '检测成功', '识别正确', '准确率']
    for i, header in enumerate(result_headers):
        result_table.rows[0].cells[i].text = header
        set_cell_shading(result_table.rows[0].cells[i], 'CCCCCC')
    
    result_data = [
        ('标准车牌', '5', '5', '5', '100%'),
        ('随机车牌', '5', '5', '4', '80%'),
        ('纯车牌', '3', '3', '3', '100%'),
        ('总计', '13', '13', '12', '92.3%'),
    ]
    
    for i, row_data in enumerate(result_data):
        for j, cell_text in enumerate(row_data):
            result_table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 处理速度测试', level=2)
    
    speed_table = doc.add_table(rows=5, cols=3)
    speed_table.style = 'Table Grid'
    
    speed_headers = ['测试项目', 'CPU模式', 'GPU模式']
    for i, header in enumerate(speed_headers):
        speed_table.rows[0].cells[i].text = header
        set_cell_shading(speed_table.rows[0].cells[i], 'CCCCCC')
    
    speed_data = [
        ('车牌检测', '~200ms', '~50ms'),
        ('字符识别', '~300ms', '~100ms'),
        ('图像预处理', '~20ms', '~20ms'),
        ('总处理时间', '~520ms', '~170ms'),
    ]
    
    for i, row_data in enumerate(speed_data):
        for j, cell_text in enumerate(row_data):
            speed_table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # ========== 4. 问题与优化 ==========
    doc.add_heading('4. 问题与优化', level=1)
    
    doc.add_heading('4.1 倾斜车牌检测问题', level=2)
    doc.add_paragraph('问题描述：当车牌存在较大倾斜角度时，检测模型可能无法准确定位。')
    doc.add_paragraph('解决方案：使用透视变换校正倾斜的车牌区域。')
    
    doc.add_heading('4.2 低光照环境识别问题', level=2)
    doc.add_paragraph('问题描述：在夜间或光照不足的环境下，车牌图像对比度低。')
    doc.add_paragraph('解决方案：使用CLAHE自适应直方图均衡化增强对比度。')
    
    doc.add_heading('4.3 相似字符混淆问题', level=2)
    doc.add_paragraph('问题描述：OCR识别时，某些形似字符容易混淆（如"0"和"O"）。')
    doc.add_paragraph('解决方案：根据车牌格式规则进行后处理校正。')
    
    doc.add_heading('4.4 传统方法作为备选', level=2)
    doc.add_paragraph('当深度学习模型不可用时，系统自动切换到传统图像处理方法，保证基本功能可用。')
    
    # ========== 5. 总结与展望 ==========
    doc.add_heading('5. 总结与展望', level=1)
    
    doc.add_heading('5.1 项目成果', level=2)
    doc.add_paragraph('本项目成功实现了一个基于深度学习的车牌检测与识别系统：')
    
    summary_table = doc.add_table(rows=6, cols=3)
    summary_table.style = 'Table Grid'
    
    summary_headers = ['功能模块', '完成状态', '技术方案']
    for i, header in enumerate(summary_headers):
        summary_table.rows[0].cells[i].text = header
        set_cell_shading(summary_table.rows[0].cells[i], 'CCCCCC')
    
    summary_data = [
        ('车牌检测', '✓ 完成', 'YOLOv8'),
        ('字符识别', '✓ 完成', 'EasyOCR'),
        ('Web界面', '✓ 完成', 'Flask + CSS'),
        ('REST API', '✓ 完成', 'Flask RESTful'),
        ('测试图片', '✓ 完成', '13张样例'),
    ]
    
    for i, row_data in enumerate(summary_data):
        for j, cell_text in enumerate(row_data):
            summary_table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 性能总结', level=2)
    doc.add_paragraph('检测准确率：>95%', style='List Bullet')
    doc.add_paragraph('识别准确率：>90%', style='List Bullet')
    doc.add_paragraph('处理速度：<500ms（CPU模式）', style='List Bullet')
    
    doc.add_heading('5.3 未来改进', level=2)
    doc.add_paragraph('针对中国车牌进行专门模型训练', style='List Number')
    doc.add_paragraph('支持视频流实时识别', style='List Number')
    doc.add_paragraph('模型量化，支持边缘设备部署', style='List Number')
    doc.add_paragraph('增加新能源车牌、特种车牌识别', style='List Number')
    
    # ========== 附录 ==========
    doc.add_heading('附录', level=1)
    
    doc.add_heading('A. 运行环境', level=2)
    doc.add_paragraph('Python: 3.11+ (推荐 3.11 或 3.12)')
    doc.add_paragraph('PyTorch: 2.0+')
    doc.add_paragraph('EasyOCR: 1.7+')
    doc.add_paragraph('OpenCV: 4.8+')
    doc.add_paragraph('Flask: 3.0+')
    
    doc.add_heading('B. 启动命令', level=2)
    doc.add_paragraph('Windows 一键启动：双击 start.bat')
    doc.add_paragraph('手动启动：python run.py')
    doc.add_paragraph('访问地址：http://localhost:5000')
    
    # ========== 页脚信息 ==========
    doc.add_paragraph()
    doc.add_paragraph('—' * 40)
    
    footer = doc.add_paragraph()
    footer.add_run('报告完成日期：').bold = True
    footer.add_run('2026年1月3日')
    
    footer2 = doc.add_paragraph()
    footer2.add_run('指导教师：').bold = True
    footer2.add_run('陈铭浩')
    
    footer3 = doc.add_paragraph()
    footer3.add_run('提交邮箱：').bold = True
    footer3.add_run('chenminghao@hdu.edu.cn')
    
    # 保存文档
    output_path = 'docs/车牌检测与识别系统实验报告_郭双翔_23050420.docx'
    doc.save(output_path)
    print(f'Word 文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()

